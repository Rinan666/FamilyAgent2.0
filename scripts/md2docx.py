#!/usr/bin/env python3
"""
Markdown → Word 转换脚本

将商业计划书和项目开发计划书转为专业排版的 .docx 文件

用法：
    python scripts/md2docx.py
"""
import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================
# 配置
# ============================================

PROJECT_ROOT = Path(__file__).parent.parent

FILES = [
    {
        "input": PROJECT_ROOT / "商业计划书.md",
        "output": PROJECT_ROOT / "docs" / "家族教育Agent-商业计划书.docx",
        "title": "家族教育Agent 商业计划书",
    },
    {
        "input": PROJECT_ROOT / "项目开发计划书.md",
        "output": PROJECT_ROOT / "docs" / "家族教育Agent-项目开发计划书.docx",
        "title": "家族教育Agent 项目开发计划书",
    },
]

# 中文友好字体
FONT_BODY = "微软雅黑"
FONT_HEADING = "微软雅黑"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11.5)
FONT_SIZE_CODE = Pt(9)
FONT_SIZE_TABLE = Pt(10)
FONT_SIZE_FOOTER = Pt(9)

# 颜色
COLOR_HEADING = RGBColor(0x1A, 0x56, 0xDB)  # 深蓝
COLOR_CODE_BG = "F5F5F5"
COLOR_TABLE_HEADER = "1A56DB"
COLOR_TABLE_HEADER_TEXT = "FFFFFF"
COLOR_TABLE_STRIPE = "F8FAFC"
COLOR_BORDER = "D1D5DB"


# ============================================
# 样式设置
# ============================================

def setup_styles(doc: Document):
    """配置文档样式"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = FONT_SIZE_BODY
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35

    # 标题样式
    for level, (size, color) in enumerate(
        [(FONT_SIZE_H1, COLOR_HEADING), (FONT_SIZE_H2, COLOR_HEADING),
         (FONT_SIZE_H3, COLOR_HEADING), (FONT_SIZE_H4, RGBColor(0x37, 0x41, 0x51))],
        start=1,
    ):
        heading_style = doc.styles[f"Heading {level}"]
        hf = heading_style.font
        hf.name = FONT_HEADING
        hf.size = size
        hf.color.rgb = color
        hf.bold = True
        heading_style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        hp = heading_style.paragraph_format
        hp.space_before = Pt(18 if level <= 2 else 12)
        hp.space_after = Pt(8 if level <= 2 else 6)

    # 代码样式
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    cf = code_style.font
    cf.name = FONT_CODE
    cf.size = FONT_SIZE_CODE
    cf.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    cp = code_style.paragraph_format
    cp.space_before = Pt(0)
    cp.space_after = Pt(0)
    cp.line_spacing = 1.2

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)


# ============================================
# Markdown 解析
# ============================================

def parse_markdown(text: str) -> list[dict]:
    """解析Markdown为结构化元素列表"""
    lines = text.split("\n")
    elements = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # 水平线
        if re.match(r"^---+$", line.strip()):
            elements.append({"type": "hr"})
            i += 1
            continue

        # 代码块
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            elements.append({"type": "code", "content": "\n".join(code_lines)})
            continue

        # 表格
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|?$", lines[i + 1].strip()):
            table_data = parse_table(lines, i)
            elements.append({"type": "table", "headers": table_data["headers"], "rows": table_data["rows"]})
            i = table_data["end"]
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text_content = clean_inline(heading_match.group(2))
            # 跳过目录链接中的标题（含有 [] 的）
            if not re.match(r"^\[.+\]\(.+\)$", heading_match.group(2).strip()):
                elements.append({"type": "heading", "level": level, "content": text_content})
            i += 1
            continue

        # 无序列表
        list_match = re.match(r"^(\s*)[\-*]\s+(.+)$", line)
        if list_match:
            items = []
            indent = len(list_match.group(1))
            while i < len(lines):
                lm = re.match(r"^(\s*)[\-*]\s+(.+)$", lines[i])
                if not lm:
                    break
                items.append(clean_inline(lm.group(2)))
                i += 1
            elements.append({"type": "list", "items": items, "ordered": False})
            continue

        # 有序列表
        num_list_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if num_list_match:
            items = []
            while i < len(lines):
                nlm = re.match(r"^(\s*)\d+\.\s+(.+)$", lines[i])
                if not nlm:
                    break
                items.append(clean_inline(nlm.group(2)))
                i += 1
            elements.append({"type": "list", "items": items, "ordered": True})
            continue

        # 引用
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            elements.append({"type": "quote", "content": " ".join(quote_lines)})
            continue

        # 空行
        if line.strip() == "":
            elements.append({"type": "blank"})
            i += 1
            continue

        # 普通段落
        para_lines = []
        while i < len(lines) and lines[i].strip() != "" and not lines[i].strip().startswith("```") \
                and not re.match(r"^(#{1,6})\s+", lines[i]) \
                and not ("|" in lines[i] and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|?$", lines[i + 1].strip())) \
                and not re.match(r"^(\s*)[\-*]\s+", lines[i]) \
                and not re.match(r"^(\s*)\d+\.\s+", lines[i]) \
                and not re.match(r"^---+$", lines[i].strip()) \
                and not lines[i].strip().startswith(">"):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            text_content = clean_inline(" ".join(para_lines))
            # 跳过纯目录链接行
            if text_content.strip() and not re.match(r"^[\d\.\s]+\[.+\]\(.+\)[\s\d\.\[\]\(\)\#\-]*$", text_content.strip()):
                elements.append({"type": "paragraph", "content": text_content})
        else:
            i += 1

    return elements


def parse_table(lines: list[str], start: int) -> dict:
    """解析Markdown表格"""
    # header
    headers = [cell.strip() for cell in lines[start].split("|") if cell.strip() != ""]
    # skip separator line
    rows = []
    i = start + 2
    while i < len(lines) and "|" in lines[i]:
        cells = [cell.strip() for cell in lines[i].split("|") if cell.strip() != ""]
        if cells:
            rows.append(cells)
        i += 1
    return {"headers": headers, "rows": rows, "end": i}


def clean_inline(text: str) -> str:
    """清理行内格式标记，保留纯文本（后续做富文本处理）"""
    # 保留加粗标记给后续富文本处理
    return text.strip()


# ============================================
# Word 生成
# ============================================

def add_heading_element(doc: Document, level: int, text: str):
    """添加标题"""
    # 先用inline解析处理富文本
    p = doc.add_heading(level=level)
    add_inline_runs(p, text)


def add_paragraph_element(doc: Document, text: str):
    """添加段落（支持粗体等行内格式）"""
    p = doc.add_paragraph()
    add_inline_runs(p, text)


def add_inline_runs(paragraph, text: str):
    """解析行内Markdown并添加Run"""
    # 处理 **粗体**、*斜体*、`代码`
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last_end = 0

    for match in pattern.finditer(text):
        # 前面的普通文本
        if match.start() > last_end:
            normal = text[last_end:match.start()]
            run = paragraph.add_run(normal)
            run.font.name = FONT_BODY
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = FONT_BODY
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.name = FONT_BODY
            run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = FONT_CODE
            run.font.size = FONT_SIZE_CODE
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        last_end = match.end()

    # 尾部普通文本
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = FONT_BODY
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def add_code_block(doc: Document, code: str):
    """添加代码块"""
    for line_no, code_line in enumerate(code.split("\n")):
        p = doc.add_paragraph(style="CodeBlock")
        # 灰色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_CODE_BG}" w:val="clear"/>')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(code_line if code_line else " ")
        run.font.name = FONT_CODE
        run.font.size = FONT_SIZE_CODE


def add_table_element(doc: Document, headers: list[str], rows: list[list[str]]):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    for j, header in enumerate(headers):
        cell = header_cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(clean_inline(header))
        run.bold = True
        run.font.size = FONT_SIZE_TABLE
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_BODY
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        # 蓝色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row in enumerate(rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            if j < len(headers):
                cell = row_cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                # 条纹背景
                if i % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_STRIPE}" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                add_inline_runs(p, clean_inline(cell_text))
                for run in p.runs:
                    run.font.size = FONT_SIZE_TABLE
                    if not run.font.name:
                        run.font.name = FONT_BODY
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    doc.add_paragraph("")  # 表后空行


def add_list_element(doc: Document, items: list[str], ordered: bool = False):
    """添加列表"""
    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)
        prefix = f"{idx + 1}. " if ordered else "• "
        run = p.add_run(prefix)
        run.font.name = FONT_BODY
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        add_inline_runs(p, item)


def add_quote_element(doc: Document, text: str):
    """添加引用"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    # 左边框效果：缩进 + 灰色文字
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.size = Pt(10)
    run.italic = True


def add_hr_element(doc: Document):
    """添加水平分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 用底部边框模拟
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{COLOR_BORDER}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def add_cover_page(doc: Document, title: str, subtitle: str = ""):
    """添加封面页"""
    # 空行
    for _ in range(6):
        doc.add_paragraph("")

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_HEADING
    run.bold = True
    run.font.name = FONT_HEADING
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)

    # 副标题
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run2.font.name = FONT_BODY
        run2.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    doc.add_paragraph("")

    # 元信息
    info_lines = ["版本：v1.0", "日期：2026年6月", "密级：内部资料"]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run.font.name = FONT_BODY
        run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)

    # 分页
    doc.add_page_break()


# ============================================
# 主流程
# ============================================

def convert_md_to_docx(input_path: Path, output_path: Path, title: str):
    """转换单个Markdown文件为Word"""
    print(f"[*] Converting: {input_path.name} -> {output_path.name}")

    with open(input_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # 移除YAML frontmatter如果有
    md_text = re.sub(r"^---\n.*?\n---\n", "", md_text, flags=re.DOTALL)

    # 移除目录部分（在正文中效果不好）
    # 找到第一个 ## 开始处理
    elements = parse_markdown(md_text)

    # 创建Word文档
    doc = Document()
    setup_styles(doc)

    # 封面
    add_cover_page(doc, title, "家族教育Agent 系列文档")

    # 正文
    for elem in elements:
        if elem["type"] == "heading":
            add_heading_element(doc, elem["level"], elem["content"])
        elif elem["type"] == "paragraph":
            add_paragraph_element(doc, elem["content"])
        elif elem["type"] == "code":
            add_code_block(doc, elem["content"])
        elif elem["type"] == "table":
            add_table_element(doc, elem["headers"], elem["rows"])
        elif elem["type"] == "list":
            add_list_element(doc, elem["items"], elem.get("ordered", False))
        elif elem["type"] == "quote":
            add_quote_element(doc, elem["content"])
        elif elem["type"] == "hr":
            add_hr_element(doc)
        elif elem["type"] == "blank":
            # 跳过连续空行
            pass

    # 页脚
    add_hr_element(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为商业计划初稿，数据为预估值，实际执行中需根据市场反馈持续迭代 —")
    run.font.size = FONT_SIZE_FOOTER
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run.italic = True

    # 保存
    os.makedirs(output_path.parent, exist_ok=True)
    doc.save(str(output_path))
    print(f"  [OK] Saved: {output_path}")


def main():
    print("=" * 50)
    print("  Markdown -> Word Converter")
    print("=" * 50)
    print()

    for file_info in FILES:
        if not file_info["input"].exists():
            print(f"  [WARN] File not found: {file_info['input']}")
            continue
        convert_md_to_docx(
            input_path=file_info["input"],
            output_path=file_info["output"],
            title=file_info["title"],
        )

    print()
    print("All done!")
    print(f"Output: {PROJECT_ROOT / 'docs'}")


if __name__ == "__main__":
    main()
